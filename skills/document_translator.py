from docx import Document
from ai.chat_memory import ChatMemory
from skills.gmp_translator import translate_gmp_text


class DocumentTranslator:
    def __init__(self, provider='deepseek'):
        self.memory = ChatMemory(provider=provider)

    def _translate_text(self, text):
        text = text.strip()
        if not text:
            return text

        translated = self.memory.ask(
            f'Translate the following pharmaceutical/GMP text into Russian and preserve meaning, numbering, units and formatting: {text}'
        )

        return translate_gmp_text(translated)

    def translate_docx(self, input_file, output_file):
        doc = Document(input_file)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph.text = self._translate_text(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            paragraph.text = self._translate_text(paragraph.text)

        doc.save(output_file)
        return output_file
